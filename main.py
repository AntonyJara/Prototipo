from flask import Flask, request, jsonify, send_file, render_template
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openai
import os
from io import BytesIO

app = Flask(__name__)

# Configurar OpenAI API
openai.api_key = os.environ.get('OPENAI_API_KEY')

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/generar', methods=['POST'])
def generar_documento():
    try:
        data = request.json
        ciclo = data.get('ciclo')
        area = data.get('area')
        tema = data.get('tema')

        if not all([ciclo, area, tema]):
            return jsonify({'error': 'Faltan datos requeridos'}), 400

        # Generar contenido con IA
        contenido = generar_contenido_ia(ciclo, area, tema)

        # Crear documento Word
        doc = crear_documento_word(ciclo, area, tema, contenido)

        # Guardar en memoria
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        return send_file(
            file_stream,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'Competencias_{area}_Ciclo_{ciclo}.docx'
        )

    except Exception as e:
        return jsonify({'error': str(e)}), 500

def generar_contenido_ia(ciclo, area, tema):
    prompt = f'''Eres un asistente pedagógico experto en el Currículo Nacional de Educación Básica del Perú.

Genera información educativa para:
- Ciclo: {ciclo}
- Área: {area}
- Tema: {tema}

Proporciona la siguiente información en formato JSON:
{{
  "competencia": "Nombre completo de la competencia del área según el Currículo Nacional",
  "capacidades": ["Capacidad 1", "Capacidad 2", "Capacidad 3"],
  "estandar": "Estándar de aprendizaje para el ciclo {ciclo} relacionado con esta competencia",
  "criterios": ["Criterio de evaluación 1", "Criterio de evaluación 2"],
  "instrumento": "Lista de cotejo",
  "competencia_transversal": "Nombre de la competencia transversal relacionada",
  "enfoque_transversal": "Nombre del enfoque transversal",
  "descripcion_enfoque": "Breve descripción del enfoque transversal"
}}

Asegúrate de que sea información precisa según el Currículo Nacional del Perú.'''

    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "Eres un experto en educación y currículo nacional peruano."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=1500
        )

        import json
        contenido_texto = response.choices[0].message.content
        # Extraer JSON del texto
        inicio = contenido_texto.find('{')
        fin = contenido_texto.rfind('}') + 1
        contenido_json = contenido_texto[inicio:fin]

        return json.loads(contenido_json)

    except Exception as e:
        print(f"Error en IA: {e}")
        # Retornar datos de ejemplo en caso de error
        return {
            "competencia": f"Competencia principal del área de {area}",
            "capacidades": ["Capacidad relacionada 1", "Capacidad relacionada 2"],
            "estandar": f"Estándar de aprendizaje para ciclo {ciclo}",
            "criterios": ["Criterio de evaluación 1", "Criterio de evaluación 2"],
            "instrumento": "Lista de cotejo",
            "competencia_transversal": "Gestiona su aprendizaje de manera autónoma",
            "enfoque_transversal": "Enfoque de orientación al bien común",
            "descripcion_enfoque": "Promueve valores y actitudes para el bienestar colectivo"
        }

def crear_documento_word(ciclo, area, tema, contenido):
    doc = Document()

    # Configurar márgenes
    section = doc.sections[0]
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    # Título
    titulo = doc.add_heading('COMPETENCIAS Y CRITERIOS DE EVALUACIÓN', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Información básica
    doc.add_paragraph()
    info = doc.add_paragraph()
    info.add_run(f'Ciclo: ').bold = True
    info.add_run(f'{ciclo}\n')
    info.add_run(f'Área: ').bold = True
    info.add_run(f'{area}\n')
    info.add_run(f'Tema: ').bold = True
    info.add_run(f'{tema}')

    doc.add_paragraph()

    # Crear tabla
    tabla = doc.add_table(rows=1, cols=5)
    tabla.style = 'Table Grid'

    # Encabezados
    headers = ['Competencias y Capacidades', 'Estándar de Aprendizaje', 
               'Criterios de Evaluación', 'Instrumento de Evaluación', 
               'Competencia Transversal']

    for i, header in enumerate(headers):
        cell = tabla.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)

    # Agregar contenido
    row = tabla.add_row()

    # Competencias y capacidades
    comp_text = f"{contenido['competencia']}\n\n"
    comp_text += "Capacidades:\n" + "\n".join([f"• {cap}" for cap in contenido['capacidades']])
    row.cells[0].text = comp_text

    # Estándar
    row.cells[1].text = contenido['estandar']

    # Criterios
    crit_text = "\n".join([f"• {crit}" for crit in contenido['criterios']])
    row.cells[2].text = crit_text

    # Instrumento
    row.cells[3].text = contenido['instrumento']

    # Competencia transversal
    trans_text = f"{contenido['competencia_transversal']}\n\n"
    trans_text += f"Enfoque: {contenido['enfoque_transversal']}\n\n"
    trans_text += contenido['descripcion_enfoque']
    row.cells[4].text = trans_text

    # Ajustar anchos de columna
    for row in tabla.rows:
        for cell in row.cells:
            cell.width = Inches(2.0)

    # Pie de página
    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph('Documento generado por Asistente Pedagógico IA')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    return doc

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
