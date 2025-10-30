from flask import Flask, request, jsonify, send_file, render_template_string
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openai
import os
from io import BytesIO
import json

app = Flask(__name__)

# Configurar OpenAI API
openai.api_key = os.environ.get('OPENAI_API_KEY')

# BASE DE DATOS CON COMPETENCIAS REALES DEL CURRÍCULO NACIONAL PERUANO
COMPETENCIAS_DB = {
    "Comunicación": {
        "III": {
            "temas": ["Textos narrativos", "Fábulas y leyendas", "Descripción de personas y objetos", "Lectura comprensiva", "Producción de textos"],
            "competencia": "Se comunica oralmente en su lengua materna.",
            "capacidades": [
                "Obtiene información del texto oral",
                "Infiere e interpreta información del texto",
                "Adecúa, organiza y desarrolla ideas de forma coherente y cohesionada"
            ],
            "estandar": "Se comunica oralmente mediante diversos tipos de textos; identifica información explícita; realiza inferencias sencillas; se expresa con pronunciación y entonación adecuadas.",
            "criterios": [
                "Comprende información explícita en textos orales simples sobre temas cotidianos",
                "Expresa sus ideas con claridad, usando palabras apropiadas y gestos"
            ],
            "instrumento": "Lista de cotejo",
            "competencia_transversal": "Gestiona su aprendizaje de manera autónoma",
            "enfoque_transversal": "Orientación al bien común",
            "descripcion_enfoque": "Valora y practica la justicia, solidaridad y responsabilidad en la comunidad."
        },
        "IV": {
            "temas": ["Narrativa clásica", "Poesía y rima", "Cartas y mensajes", "Cuento literario", "Reportaje"],
            "competencia": "Lee diversos tipos de textos escritos en su lengua materna.",
            "capacidades": [
                "Obtiene información explícita del texto",
                "Realiza inferencias e interpreta el significado del texto",
                "Reflexiona y evalúa la forma, el contenido y contexto del texto"
            ],
            "estandar": "Lee diversos tipos de textos con estructuras complejas, vocabulario variado; identifica información; realiza inferencias; opina sobre lo leído.",
            "criterios": [
                "Identifica la información central y detalles en textos con estructura narrativa",
                "Formula opiniones sobre la intención del autor basándose en elementos textuales"
            ],
            "instrumento": "Lista de cotejo",
            "competencia_transversal": "Se desenvuelve en entornos virtuales generados por las TIC",
            "enfoque_transversal": "Inclusión o atención a la diversidad",
            "descripcion_enfoque": "Respeta y valora la diversidad de las personas en todos sus aspectos."
        },
        "V": {
            "temas": ["Literatura infantil", "Textos informativos complejos", "Ensayos cortos", "Drama y teatro", "Análisis de medios de comunicación"],
            "competencia": "Escribe diversos tipos de textos en su lengua materna.",
            "capacidades": [
                "Adecúa el texto a la situación comunicativa",
                "Organiza y desarrolla ideas de forma coherente y cohesionada",
                "Utiliza convenciones del lenguaje escrito de forma apropiada"
            ],
            "estandar": "Escribe textos complejos de distintos tipos; adecúa su contenido a la audiencia; organiza sus ideas; usa vocabulario variado; aplica convenciones del lenguaje escrito.",
            "criterios": [
                "Produce textos con estructura clara, coherencia entre ideas y vocabulario pertinente",
                "Revisa y corrige sus textos considerando puntuación, ortografía y coherencia"
            ],
            "instrumento": "Lista de cotejo",
            "competencia_transversal": "Gestiona su aprendizaje de manera autónoma",
            "enfoque_transversal": "Derechos humanos",
            "descripcion_enfoque": "Promueve y defiende los derechos humanos para todos."
        }
    },
    "Matemática": {
        "III": {
            "temas": ["Números naturales hasta 100", "Suma y resta", "Figuras geométricas", "Medición de longitud", "Patrones numéricos"],
            "competencia": "Resuelve problemas de cantidad.",
            "capacidades": [
                "Traduce cantidades a expresiones numéricas",
                "Comunica su comprensión sobre los números",
                "Usa estrategias y procedimientos de estimación y cálculo"
            ],
            "estandar": "Resuelve problemas de cantidad relacionados con agregar, quitar e igualar con números naturales hasta 100; usa estrategias de cálculo; explica su procedimiento.",
            "criterios": [
                "Resuelve sumas y restas con números hasta 100 de forma correcta",
                "Explica sus estrategias usando lenguaje matemático apropiado"
            ],
            "instrumento": "Lista de cotejo",
            "competencia_transversal": "Gestiona su aprendizaje de manera autónoma",
            "enfoque_transversal": "Búsqueda de la excelencia",
            "descripcion_enfoque": "Persigue mejorar constantemente en la realización de tareas y comprensión de conceptos."
        },
        "IV": {
            "temas": ["Fracciones", "Multiplicación y división", "Números decimales", "Áreas y perímetros", "Datos y gráficos"],
            "competencia": "Resuelve problemas de regularidad, equivalencia y cambio.",
            "capacidades": [
                "Traduce datos y condiciones a expresiones algebraicas",
                "Comunica su comprensión sobre relaciones algebraicas",
                "Usa estrategias para resolver ecuaciones"
            ],
            "estandar": "Resuelve problemas con fracciones, decimales y operaciones; identifica patrones en secuencias; usa estrategias variadas de cálculo; justifica sus procedimientos.",
            "criterios": [
                "Realiza operaciones con fracciones y decimales demostrando comprensión",
                "Reconoce y completa patrones numéricos justificando la regla"
            ],
            "instrumento": "Lista de cotejo",
            "competencia_transversal": "Se desenvuelve en entornos virtuales generados por las TIC",
            "enfoque_transversal": "Orientación al bien común",
            "descripcion_enfoque": "Busca resultados que beneficien a todos los integrantes de la comunidad."
        },
        "V": {
            "temas": ["Proporcionalidad", "Potencias y raíces", "Ecuaciones lineales", "Volumen y capacidad", "Probabilidad básica"],
            "competencia": "Resuelve problemas de forma, movimiento y localización.",
            "capacidades": [
                "Modela objetos con formas geométricas",
                "Comunica su comprensión sobre formas y relaciones geométricas",
                "Usa estrategias para medir y calcular propiedades de figuras"
            ],
            "estandar": "Resuelve problemas geométricos identificando propiedades de figuras; calcula perímetros, áreas y volúmenes; justifica sus estrategias con argumentos geométricos.",
            "criterios": [
                "Calcula correctamente áreas y perímetros de figuras compuestas",
                "Justifica sus cálculos usando propiedades geométricas"
            ],
            "instrumento": "Lista de cotejo",
            "competencia_transversal": "Gestiona su aprendizaje de manera autónoma",
            "enfoque_transversal": "Inclusión o atención a la diversidad",
            "descripcion_enfoque": "Valora la diversidad de estrategias de resolución de problemas."
        }
    },
    "Personal Social": {
        "III": {
            "temas": ["Mi familia y comunidad", "Normas en el hogar", "Costumbres y tradiciones", "El respeto", "Seguridad personal"],
            "competencia": "Construye su identidad.",
            "capacidades": [
                "Se valora a sí mismo",
                "Autorregula sus emociones",
                "Reflexiona sobre sus prácticas culturales"
            ],
            "estandar": "Conoce sus características personales; identifica sus emociones y las de otros; respeta las diferencias; practica valores como respeto y responsabilidad.",
            "criterios": [
                "Expresa sus emociones e identifica cómo se sienten los demás",
                "Sigue normas de convivencia en diferentes contextos"
            ],
            "instrumento": "Lista de cotejo",
            "competencia_transversal": "Gestiona su aprendizaje de manera autónoma",
            "enfoque_transversal": "Derechos humanos",
            "descripcion_enfoque": "Promueve el reconocimiento y respeto de los derechos de todos."
        },
        "IV": {
            "temas": ["Organización social", "Roles familiares", "Tradiciones regionales", "Convivencia democrática", "Responsabilidad ciudadana"],
            "competencia": "Convive y participa democráticamente.",
            "capacidades": [
                "Interactúa con todas las personas",
                "Construye normas y asume acuerdos y leyes",
                "Participa en asuntos públicos"
            ],
            "estandar": "Practica la empatía; construye normas consensuadas; participa en decisiones comunitarias; muestra disposición a trabajar en equipo.",
            "criterios": [
                "Propone soluciones pacíficas ante conflictos cotidianos",
                "Participa activamente en la toma de decisiones del grupo"
            ],
            "instrumento": "Lista de cotejo",
            "competencia_transversal": "Gestiona su aprendizaje de manera autónoma",
            "enfoque_transversal": "Orientación al bien común",
            "descripcion_enfoque": "Trabaja por el bienestar común considerando los intereses de todos."
        },
        "V": {
            "temas": ["Ciudadanía y participación", "Instituciones públicas", "Derechos y deberes", "Patrimonio cultural", "Sostenibilidad ambiental"],
            "competencia": "Gestiona responsablemente el espacio y el ambiente.",
            "capacidades": [
                "Comprende las dinámicas entre elementos naturales y sociales",
                "Maneja responsablemente recursos",
                "Evalúa problemáticas ambientales"
            ],
            "estandar": "Identifica elementos naturales y sociales; evalúa problemáticas ambientales; propone soluciones sostenibles; reconoce su rol en el cuidado del ambiente.",
            "criterios": [
                "Explica cómo los elementos naturales influyen en la vida de las personas",
                "Propone acciones para conservar y proteger el ambiente"
            ],
            "instrumento": "Lista de cotejo",
            "competencia_transversal": "Se desenvuelve en entornos virtuales generados por las TIC",
            "enfoque_transversal": "Inclusión o atención a la diversidad",
            "descripcion_enfoque": "Considera a todos los seres humanos y al ambiente en sus decisiones."
        }
    },
    "Ciencia y Tecnología": {
        "III": {
            "temas": ["Seres vivos y su hábitat", "El cuerpo humano", "Ciclo del agua", "Fuentes de luz", "Seguridad en el hogar"],
            "competencia": "Indaga mediante métodos científicos para construir conocimientos.",
            "capacidades": [
                "Problematiza situaciones",
                "Diseña estrategias para indagación",
                "Genera y registra datos",
                "Analiza datos e información"
            ],
            "estandar": "Realiza indagaciones simples; observa características de objetos; predice cambios; comunica sus conclusiones; plantea preguntas sobre fenómenos naturales.",
            "criterios": [
                "Realiza observaciones ordenadas de fenómenos naturales",
                "Comunica sus predicciones y conclusiones de forma clara"
            ],
            "instrumento": "Lista de cotejo",
            "competencia_transversal": "Gestiona su aprendizaje de manera autónoma",
            "enfoque_transversal": "Búsqueda de la excelencia",
            "descripcion_enfoque": "Persigue comprender los fenómenos con precisión y rigor."
        },
        "IV": {
            "temas": ["Cadenas alimenticias", "Reproducción de plantas y animales", "Estados de la materia", "Fuerzas y movimiento", "Efectos del calor"],
            "competencia": "Explica el mundo natural basándose en conocimientos científicos.",
            "capacidades": [
                "Comprende y aplica conocimientos científicos",
                "Argumenta afirmaciones sobre fenómenos naturales"
            ],
            "estandar": "Identifica características de los seres vivos; comprende procesos naturales como ciclos de vida; explica fenómenos físicos con lenguaje científico.",
            "criterios": [
                "Describe correctamente los ciclos de vida de organismos",
                "Explica relaciones causa-efecto en fenómenos naturales"
            ],
            "instrumento": "Lista de cotejo",
            "competencia_transversal": "Se desenvuelve en entornos virtuales generados por las TIC",
            "enfoque_transversal": "Orientación al bien común",
            "descripcion_enfoque": "Reconoce la importancia de la naturaleza para la vida."
        },
        "V": {
            "temas": ["Ecología y ecosistemas", "Genética básica", "Energía y recursos", "Tecnología y sostenibilidad", "Cambio climático"],
            "competencia": "Diseña y construye soluciones tecnológicas.",
            "capacidades": [
                "Determina una alternativa de solución tecnológica",
                "Diseña la alternativa de solución tecnológica",
                "Implementa y valida la alternativa"
            ],
            "estandar": "Diseña soluciones a problemas tecnológicos; valida su funcionamiento; evalúa el impacto ambiental; propone mejoras sostenibles.",
            "criterios": [
                "Diseña prototipos simples con materiales reciclables",
                "Evalúa la efectividad y sostenibilidad de su solución"
            ],
            "instrumento": "Lista de cotejo",
            "competencia_transversal": "Gestiona su aprendizaje de manera autónoma",
            "enfoque_transversal": "Búsqueda de la excelencia",
            "descripcion_enfoque": "Busca mejorar continuamente sus diseños tecnológicos."
        }
    },
    "Educación Religiosa": {
        "III": {
            "temas": ["Dios creador", "Virtudes cristianas", "La Biblia", "Principales festividades religiosas", "Amor y fraternidad"],
            "competencia": "Construye su identidad como persona humana, amada por Dios.",
            "capacidades": [
                "Valora su dignidad personal",
                "Reconoce la obra creadora de Dios",
                "Reflexiona sobre valores religiosos"
            ],
            "estandar": "Reconoce que Dios lo ama; identifica valores como bondad, respeto y solidaridad; participa en celebraciones religiosas; respeta crencias diferentes.",
            "criterios": [
                "Expresa su fe reconociendo que Dios lo ama como persona",
                "Practica virtudes cristianas en su convivencia diaria"
            ],
            "instrumento": "Lista de cotejo",
            "competencia_transversal": "Gestiona su aprendizaje de manera autónoma",
            "enfoque_transversal": "Derechos humanos",
            "descripcion_enfoque": "Reconoce la dignidad de la persona como creada a imagen de Dios."
        },
        "IV": {
            "temas": ["Jesucristo redentor", "Mandamientos de la ley de Dios", "Sacramentos", "Comunidades religiosas", "Servicio al prójimo"],
            "competencia": "Asume la experiencia del encuentro personal y comunitario con Dios.",
            "capacidades": [
                "Se relaciona con Dios en forma auténtica",
                "Experimenta encuentros con Dios",
                "Valora la vida en comunidad"
            ],
            "estandar": "Experimenta que Dios ama al ser humano; vive valores como caridad y justicia; participa en acciones comunitarias; respeta opciones religiosas.",
            "criterios": [
                "Identifica acciones que demuestran el amor de Dios en la vida",
                "Realiza compromisos de servicio hacia los demás"
            ],
            "instrumento": "Lista de cotejo",
            "competencia_transversal": "Se desenvuelve en entornos virtuales generados por las TIC",
            "enfoque_transversal": "Orientación al bien común",
            "descripcion_enfoque": "Busca el bien común y la construcción de comunidades solidarias."
        },
        "V": {
            "temas": ["Encíclicas y doctrina social", "Ecología integral", "Justicia social", "Diálogo interreligioso", "Responsabilidad moral"],
            "competencia": "Actúa coherentemente en razón de su fe según los principios de su conciencia moral.",
            "capacidades": [
                "Practica virtudes morales",
                "Toma decisiones responsables",
                "Compromete con la justicia y paz"
            ],
            "estandar": "Actúa según sus principios morales; busca la justicia y paz; respeta otras tradiciones religiosas; trabaja por el bien común.",
            "criterios": [
                "Fundamenta sus decisiones morales basándose en su fe",
                "Participa en acciones de justicia social y paz"
            ],
            "instrumento": "Lista de cotejo",
            "competencia_transversal": "Gestiona su aprendizaje de manera autónoma",
            "enfoque_transversal": "Inclusión o atención a la diversidad",
            "descripcion_enfoque": "Valora el diálogo como medio para entender diferentes perspectivas."
        }
    },
    "Arte y Cultura": {
        "III": {
            "temas": ["Expresión artística", "Dibujo y pintura", "Canciones infantiles", "Danzas folclóricas", "Artesanía local"],
            "competencia": "Aprecia críticamente manifestaciones artístico-culturales.",
            "capacidades": [
                "Percibe manifestaciones artísticas",
                "Contextualiza manifestaciones artísticas",
                "Reflexiona creadora y críticamente"
            ],
            "estandar": "Aprecia manifestaciones artísticas; identifica elementos visuales y sonoros; comparte opiniones sobre arte; respeta expresiones culturales diferentes.",
            "criterios": [
                "Identifica elementos de color y forma en obras artísticas",
                "Expresa qué siente ante diferentes manifestaciones culturales"
            ],
            "instrumento": "Lista de cotejo",
            "competencia_transversal": "Gestiona su aprendizaje de manera autónoma",
            "enfoque_transversal": "Búsqueda de la excelencia",
            "descripcion_enfoque": "Busca apreciar el arte con sensibilidad y profundidad."
        },
        "IV": {
            "temas": ["Técnicas de pintura", "Música regional", "Teatro de títeres", "Patrimonio cultural local", "Cerámica y escultura"],
            "competencia": "Crea proyectos artísticos.",
            "capacidades": [
                "Genera ideas artísticas",
                "Planifica proyectos artísticos",
                "Ejecuta técnicas artísticas",
                "Evalúa proyectos"
            ],
            "estandar": "Crea proyectos artísticos combinando elementos visuales y sonoros; utiliza técnicas apropiadas; experimenta con diferentes materiales; reflexiona sobre su proceso.",
            "criterios": [
                "Utiliza técnicas de pintura o escultura con creatividad",
                "Explica el proceso y significado de su obra artística"
            ],
            "instrumento": "Lista de cotejo",
            "competencia_transversal": "Se desenvuelve en entornos virtuales generados por las TIC",
            "enfoque_transversal": "Orientación al bien común",
            "descripcion_enfoque": "Reconoce el arte como expresión de la identidad comunitaria."
        },
        "V": {
            "temas": ["Historia del arte", "Artes visuales contemporáneas", "Música clásica y moderna", "Danza contemporánea", "Cine y audiovisual"],
            "competencia": "Se expresa artísticamente a través de diversos lenguajes.",
            "capacidades": [
                "Explora técnicas artísticas",
                "Desarrolla ideas artísticas",
                "Utiliza materiales y herramientas",
                "Reflexiona sobre procesos artísticos"
            ],
            "estandar": "Se expresa artísticamente de forma creativa; domina técnicas variadas; comunica su visión artística; analiza y valora obras de otros.",
            "criterios": [
                "Crea obras originales usando técnicas diversas",
                "Analiza críticamente obras de arte identificando técnica y mensaje"
            ],
            "instrumento": "Lista de cotejo",
            "competencia_transversal": "Gestiona su aprendizaje de manera autónoma",
            "enfoque_transversal": "Búsqueda de la excelencia",
            "descripcion_enfoque": "Persigue la excelencia en su expresión artística."
        }
    },
    "Educación Física": {
        "III": {
            "temas": ["Juegos y movimiento corporal", "Actividades lúdicas", "Higiene y salud", "Ejercicios de coordinación", "Seguridad en el juego"],
            "competencia": "Interactúa a través de sus habilidades sociomotrices.",
            "capacidades": [
                "Produce movimientos variados",
                "Colabora en juegos",
                "Valora el trabajo en equipo"
            ],
            "estandar": "Realiza movimientos variados con control; participa en juegos respetando reglas; colabora en equipo; demuestra seguridad corporal.",
            "criterios": [
                "Ejecuta movimientos coordinados en actividades lúdicas",
                "Respeta reglas y normas en juegos colectivos"
            ],
            "instrumento": "Lista de cotejo",
            "competencia_transversal": "Gestiona su aprendizaje de manera autónoma",
            "enfoque_transversal": "Orientación al bien común",
            "descripcion_enfoque": "Trabaja en equipo por el logro de objetivos comunes."
        },
        "IV": {
            "temas": ["Deportes individuales", "Deportes de equipo", "Acondicionamiento físico", "Primeros auxilios básicos", "Nutrición y actividad física"],
            "competencia": "Asume una vida saludable.",
            "capacidades": [
                "Comprende las relaciones entre actividad, nutrición, postura y salud",
                "Incorpora prácticas saludables",
                "Mantiene su cuerpo saludable"
            ],
            "estandar": "Participa en actividades físicas con técnica apropiada; practica hábitos saludables; cuida su cuerpo; desarrolla capacidades físicas.",
            "criterios": [
                "Realiza ejercicios de acondicionamiento físico correctamente",
                "Explica la importancia de la nutrición en la actividad física"
            ],
            "instrumento": "Lista de cotejo",
            "competencia_transversal": "Se desenvuelve en entornos virtuales generados por las TIC",
            "enfoque_transversal": "Inclusión o atención a la diversidad",
            "descripcion_enfoque": "Adapta actividades respetando capacidades y limitaciones de todos."
        },
        "V": {
            "temas": ["Atletismo", "Deportes de contacto", "Danza y movimiento", "Manejo del estrés", "Vida activa y bienestar"],
            "competencia": "Se desenvuelve de manera autónoma a través de su motricidad.",
            "capacidades": [
                "Comprende su cuerpo y sus movimientos",
                "Ejecuta movimientos variados",
                "Evalúa su desempeño"
            ],
            "estandar": "Ejecuta movimientos variados con precisión y eficacia; desarrolla capacidades físicas; practica actividad física regular; evalúa su progreso.",
            "criterios": [
                "Realiza técnicas deportivas con eficacia y control",
                "Evalúa su desempeño físico e identifica áreas de mejora"
            ],
            "instrumento": "Lista de cotejo",
            "competencia_transversal": "Gestiona su aprendizaje de manera autónoma",
            "enfoque_transversal": "Búsqueda de la excelencia",
            "descripcion_enfoque": "Busca mejorar continuamente su desempeño físico."
        }
    }
}

# HTML INTEGRADO
HTML_TEMPLATE = """<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Asistente Pedagógico - Generador de Competencias</title>
    <style>
        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }

        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            min-height: 100vh;
            display: flex;
            justify-content: center;
            align-items: center;
            padding: 20px;
        }

        .container {
            background: white;
            padding: 40px;
            border-radius: 20px;
            box-shadow: 0 20px 60px rgba(0,0,0,0.3);
            max-width: 600px;
            width: 100%;
        }

        h1 {
            color: #667eea;
            text-align: center;
            margin-bottom: 10px;
            font-size: 28px;
        }

        .subtitle {
            text-align: center;
            color: #666;
            margin-bottom: 30px;
            font-size: 14px;
        }

        .form-group {
            margin-bottom: 25px;
        }

        label {
            display: block;
            margin-bottom: 8px;
            color: #333;
            font-weight: 600;
            font-size: 14px;
        }

        select, input[type="text"] {
            width: 100%;
            padding: 12px 15px;
            border: 2px solid #e0e0e0;
            border-radius: 8px;
            font-size: 15px;
            transition: all 0.3s;
            font-family: inherit;
        }

        select:focus, input[type="text"]:focus {
            outline: none;
            border-color: #667eea;
            box-shadow: 0 0 0 3px rgba(102, 126, 234, 0.1);
        }

        button {
            width: 100%;
            padding: 15px;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            border: none;
            border-radius: 8px;
            font-size: 16px;
            font-weight: 600;
            cursor: pointer;
            transition: transform 0.2s, box-shadow 0.2s;
        }

        button:hover {
            transform: translateY(-2px);
            box-shadow: 0 10px 20px rgba(102, 126, 234, 0.3);
        }

        button:active {
            transform: translateY(0);
        }

        button:disabled {
            background: #ccc;
            cursor: not-allowed;
            transform: none;
        }

        .loading {
            display: none;
            text-align: center;
            margin-top: 20px;
            color: #667eea;
            font-weight: 600;
        }

        .spinner {
            border: 3px solid #f3f3f3;
            border-top: 3px solid #667eea;
            border-radius: 50%;
            width: 40px;
            height: 40px;
            animation: spin 1s linear infinite;
            margin: 20px auto;
        }

        @keyframes spin {
            0% { transform: rotate(0deg); }
            100% { transform: rotate(360deg); }
        }

        .error {
            background: #fee;
            color: #c33;
            padding: 15px;
            border-radius: 8px;
            margin-top: 20px;
            display: none;
        }

        .success {
            background: #efe;
            color: #3c3;
            padding: 15px;
            border-radius: 8px;
            margin-top: 20px;
            display: none;
            text-align: center;
        }

        .tema-list {
            background: #f5f5f5;
            padding: 10px;
            border-radius: 5px;
            margin-top: 5px;
            font-size: 12px;
            color: #666;
            max-height: 100px;
            overflow-y: auto;
        }
    </style>
</head>
<body>
    <div class="container">
        <h1>🎓 Asistente Pedagógico</h1>
        <p class="subtitle">Generador de Competencias y Criterios de Evaluación</p>

        <form id="competenciaForm">
            <div class="form-group">
                <label for="ciclo">Ciclo Educativo:</label>
                <select id="ciclo" name="ciclo" required>
                    <option value="">Selecciona un ciclo</option>
                    <option value="III">III Ciclo (1° y 2° grado)</option>
                    <option value="IV">IV Ciclo (3° y 4° grado)</option>
                    <option value="V">V Ciclo (5° y 6° grado)</option>
                </select>
            </div>

            <div class="form-group">
                <label for="area">Área Curricular:</label>
                <select id="area" name="area" required>
                    <option value="">Selecciona un área</option>
                    <option value="Comunicación">Comunicación</option>
                    <option value="Matemática">Matemática</option>
                    <option value="Personal Social">Personal Social</option>
                    <option value="Ciencia y Tecnología">Ciencia y Tecnología</option>
                    <option value="Educación Religiosa">Educación Religiosa</option>
                    <option value="Arte y Cultura">Arte y Cultura</option>
                    <option value="Educación Física">Educación Física</option>
                </select>
            </div>

            <div class="form-group">
                <label for="tema">Tema de la Unidad:</label>
                <input type="text" id="tema" name="tema" placeholder="Escribe el tema o selecciona uno sugerido" required>
                <div class="tema-list" id="temaList"></div>
            </div>

            <button type="submit" id="submitBtn">Generar Documento</button>
        </form>

        <div class="loading" id="loading">
            <div class="spinner"></div>
            <p>Generando documento con IA, por favor espera...</p>
        </div>

        <div class="error" id="error"></div>
        <div class="success" id="success"></div>
    </div>

    <script>
        const temasDB = """ + json.dumps(COMPETENCIAS_DB) + """;

        document.getElementById('ciclo').addEventListener('change', actualizarTemas);
        document.getElementById('area').addEventListener('change', actualizarTemas);

        function actualizarTemas() {
            const ciclo = document.getElementById('ciclo').value;
            const area = document.getElementById('area').value;
            const temaList = document.getElementById('temaList');

            if (ciclo && area && temasDB[area] && temasDB[area][ciclo]) {
                const temas = temasDB[area][ciclo].temas;
                temaList.innerHTML = '<strong>Temas sugeridos:</strong><br>' + temas.join(' • ');
            } else {
                temaList.innerHTML = '';
            }
        }

        document.getElementById('competenciaForm').addEventListener('submit', async (e) => {
            e.preventDefault();

            const submitBtn = document.getElementById('submitBtn');
            const loading = document.getElementById('loading');
            const error = document.getElementById('error');
            const success = document.getElementById('success');

            error.style.display = 'none';
            success.style.display = 'none';

            submitBtn.disabled = true;
            loading.style.display = 'block';

            const formData = {
                ciclo: document.getElementById('ciclo').value,
                area: document.getElementById('area').value,
                tema: document.getElementById('tema').value
            };

            try {
                const response = await fetch('/generar', {
                    method: 'POST',
                    headers: {
                        'Content-Type': 'application/json'
                    },
                    body: JSON.stringify(formData)
                });

                if (response.ok) {
                    const blob = await response.blob();
                    const url = window.URL.createObjectURL(blob);
                    const a = document.createElement('a');
                    a.href = url;
                    a.download = `Competencias_${formData.area}_Ciclo_${formData.ciclo}.docx`;
                    document.body.appendChild(a);
                    a.click();
                    window.URL.revokeObjectURL(url);
                    document.body.removeChild(a);

                    success.textContent = '✅ Documento generado exitosamente. Descarga iniciada.';
                    success.style.display = 'block';
                } else {
                    const errorData = await response.json();
                    throw new Error(errorData.error || 'Error al generar el documento');
                }
            } catch (err) {
                error.textContent = '❌ ' + err.message;
                error.style.display = 'block';
            } finally {
                loading.style.display = 'none';
                submitBtn.disabled = false;
            }
        });
    </script>
</body>
</html>"""

@app.route('/')
def index():
    return render_template_string(HTML_TEMPLATE)

@app.route('/generar', methods=['POST'])
def generar_documento():
    try:
        data = request.json
        ciclo = data.get('ciclo')
        area = data.get('area')
        tema = data.get('tema')

        if not all([ciclo, area, tema]):
            return jsonify({'error': 'Faltan datos requeridos'}), 400

        # Primero intenta buscar en la base de datos local
        contenido = buscar_en_db(ciclo, area)

        # Si no encuentra en DB, intenta con IA
        if not contenido:
            contenido = generar_contenido_ia(ciclo, area, tema)

        # Crear documento Word
        doc = crear_documento_word(ciclo, area, tema, contenido)

        # Guardar en memoria
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        return send_file(
            file_stream,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'Competencias_{area}_Ciclo_{ciclo}.docx'
        )

    except Exception as e:
        return jsonify({'error': str(e)}), 500

def buscar_en_db(ciclo, area):
    try:
        if area in COMPETENCIAS_DB and ciclo in COMPETENCIAS_DB[area]:
            return COMPETENCIAS_DB[area][ciclo]
    except:
        pass
    return None

def generar_contenido_ia(ciclo, area, tema):
    prompt = f"Eres un asistente pedagógico experto en el Currículo Nacional de Educación Básica del Perú. Genera información educativa para Ciclo {ciclo}, Área {area}, Tema {tema}. Responde SOLO en JSON válido sin explicaciones. Incluye: competencia, capacidades (lista de 3), estandar, criterios (lista de 2), instrumento (siempre 'Lista de cotejo'), competencia_transversal, enfoque_transversal, descripcion_enfoque."

    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "Eres experto en educación peruana. Responde en JSON válido.  "},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=1500
        )

        contenido_texto = response.choices[0].message.content.strip()
        inicio = contenido_texto.find('{')
        fin = contenido_texto.rfind('}') + 1

        if inicio >= 0 and fin > inicio:
            return json.loads(contenido_texto[inicio:fin])
        else:
            return generar_contenido_generico(ciclo, area, tema)

    except:
        return generar_contenido_generico(ciclo, area, tema)

def generar_contenido_generico(ciclo, area, tema):
    return {
        "competencia": f"Competencia del área de {area} - Ciclo {ciclo}",
        "capacidades": [
            f"Capacidad 1: Relacionada con {tema}",
            f"Capacidad 2: Aplicación de conceptos en {tema}",
            f"Capacidad 3: Reflexión crítica sobre {tema}"
        ],
        "estandar": f"Estándar de aprendizaje para ciclo {ciclo} en el área de {area}, relacionado con {tema}",
        "criterios": [
            f"Demuestra comprensión de conceptos clave en {tema}",
            f"Aplica estrategias apropiadas para resolver situaciones relacionadas con {tema}"
        ],
        "instrumento": "Lista de cotejo",
        "competencia_transversal": "Gestiona su aprendizaje de manera autónoma",
        "enfoque_transversal": "Orientación al bien común",
        "descripcion_enfoque": "Busca el beneficio común y la construcción de comunidades solidarias."
    }

def crear_documento_word(ciclo, area, tema, contenido):
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    titulo = doc.add_heading('COMPETENCIAS Y CRITERIOS DE EVALUACIÓN', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.add_run('Ciclo: ').bold = True
    info.add_run(f'{ciclo}\n')
    info.add_run('Área: ').bold = True
    info.add_run(f'{area}\n')
    info.add_run('Tema: ').bold = True
    info.add_run(f'{tema}')

    doc.add_paragraph()

    tabla = doc.add_table(rows=1, cols=5)
    tabla.style = 'Table Grid'

    headers = ['Competencias y Capacidades', 'Estándar de Aprendizaje', 
               'Criterios de Evaluación', 'Instrumento de Evaluación', 
               'Competencia Transversal']

    for i, header in enumerate(headers):
        cell = tabla.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)

    row = tabla.add_row()

    comp_text = f"{contenido.get('competencia', 'N/A')}\n\nCapacidades:\n"
    capacidades = contenido.get('capacidades', [])
    comp_text += "\n".join([f"• {cap}" for cap in capacidades])
    row.cells[0].text = comp_text

    row.cells[1].text = contenido.get('estandar', 'No disponible')

    criterios = contenido.get('criterios', [])
    crit_text = "\n".join([f"• {crit}" for crit in criterios])
    row.cells[2].text = crit_text if crit_text else 'No disponible'

    row.cells[3].text = contenido.get('instrumento', 'No disponible')

    trans_text = f"{contenido.get('competencia_transversal', 'N/A')}\n\n"
    trans_text += f"Enfoque: {contenido.get('enfoque_transversal', 'N/A')}\n\n"
    trans_text += contenido.get('descripcion_enfoque', 'N/A')
    row.cells[4].text = trans_text

    for row in tabla.rows:
        for cell in row.cells:
            cell.width = Inches(2.0)

    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph('Documento generado por Asistente Pedagógico - Plataforma Educativa')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    return doc

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=False)
